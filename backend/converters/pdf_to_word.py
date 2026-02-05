# -*- coding: utf-8 -*-
"""
PDF to Word Converter
pdf2docx ile PDF'ten Word'e dönüştürme

NOT: pdf2docx modülü opencv gerektirdiği için LAZY IMPORT yapılıyor.
Bu sayede gunicorn boot sırasında cv2 import edilmiyor.
"""

import io
import os
import tempfile
from typing import Optional, List, TYPE_CHECKING

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.pdf_reader import PDFReader
from translators.multi_translator import get_translator
from core.font_manager import FontManager

# pdf2docx LAZY IMPORT - opencv bağımlılığı nedeniyle
# Boot sırasında import edilmiyor, sadece fonksiyon çağrıldığında
def _get_pdf2docx_converter():
    """pdf2docx Converter'ı lazy olarak import et"""
    from pdf2docx import Converter
    return Converter


class PDFToWordConverter:
    """
    PDF'ten Word'e dönüştürücü
    Görselleri ve metin düzenini korur
    """

    def __init__(self):
        self.translator = get_translator()

    def convert(self, pdf_bytes: bytes, source_lang: str = "auto",
               target_lang: str = "tr", translate: bool = False) -> bytes:
        """
        PDF'i Word'e dönüştür

        Args:
            pdf_bytes: PDF bayt verisi
            source_lang: Kaynak dil
            target_lang: Hedef dil
            translate: Çeviri yap

        Returns:
            bytes: DOCX bayt verisi
        """
        # Geçici dosyalar oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pdf_file:
            pdf_file.write(pdf_bytes)
            pdf_path = pdf_file.name

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_file:
            docx_path = docx_file.name

        try:
            # PDF'ten Word'e dönüştür (LAZY IMPORT)
            Converter = _get_pdf2docx_converter()
            cv = Converter(pdf_path)

            if translate:
                # Çeviri ile dönüştür
                self._convert_with_translation(cv, docx_path, source_lang, target_lang)
            else:
                # Sadece dönüştür
                cv.convert(docx_path)
                cv.close()

            # Sonucu oku
            with open(docx_path, "rb") as f:
                result = f.read()

            return result

        finally:
            # Geçici dosyaları sil
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
            if os.path.exists(docx_path):
                os.unlink(docx_path)

    def _convert_with_translation(self, converter, output_path: str,
                                 source_lang: str, target_lang: str):
        """Çeviri ile Word'e dönüştür"""
        # Önce PDF'i oku
        with PDFReader(pdf_bytes=None, pdf_path=converter.pdf_file.name) as reader:
            doc = Document()

            # Her sayfa için
            for page_num in range(len(reader)):
                layout = reader.analyze_page_layout(page_num)

                # Başlık ekle
                if page_num > 0:
                    doc.add_page_break()

                title = doc.add_paragraph(f"Sayfa {page_num + 1}")
                title.runs[0].bold = True
                title.runs[0].font.size = Pt(14)

                # Metin bloklarını ekle
                for block in layout.text_blocks:
                    # Çeviri yap
                    result = self.translator.translate(
                        block.text,
                        target_lang=target_lang,
                        source_lang=source_lang
                    )

                    text = result.text if result.success else block.text

                    # Paragraf ekle
                    para = doc.add_paragraph(text)

                    # Font ayarla
                    para.runs[0].font.name = "Calibri"
                    para.runs[0].font.size = Pt(block.font_size)

                    if block.is_bold:
                        para.runs[0].bold = True
                    if block.is_italic:
                        para.runs[0].italic = True

                    # Hizalama
                    if block.alignment == "center":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif block.alignment == "right":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Görselleri ekle
                for img_block in layout.images:
                    try:
                        # Görseli kaydet
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as img_file:
                            img_file.write(img_block.image_data)
                            img_temp_path = img_file.name

                        # Word'e ekle
                        doc.add_picture(img_temp_path, width=Inches(4))

                        # Temizle
                        os.unlink(img_temp_path)

                    except Exception as e:
                        print(f"⚠️ Görsel eklenemedi: {e}")

            # Kaydet
            doc.save(output_path)
            converter.close()


class AdvancedPDFToWordConverter:
    """
    Gelişmiş PDF to Word converter
    Daha fazla kontrol ve özelleştirme
    """

    def __init__(self):
        self.translator = get_translator()

    def convert_with_pdf2docx(self, pdf_bytes: bytes) -> bytes:
        """
        pdf2docx ile doğrudan dönüşüm (en iyi görsel koruma)

        Args:
            pdf_bytes: PDF bayt verisi

        Returns:
            bytes: DOCX bayt verisi
        """
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pdf_file:
            pdf_file.write(pdf_bytes)
            pdf_path = pdf_file.name

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_file:
            docx_path = docx_file.name

        try:
            # pdf2docx ile convert (LAZY IMPORT)
            Converter = _get_pdf2docx_converter()
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()

            # Sonucu oku
            with open(docx_path, "rb") as f:
                result = f.read()

            return result

        finally:
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
            if os.path.exists(docx_path):
                os.unlink(docx_path)

    def convert_and_translate(self, pdf_bytes: bytes, source_lang: str = "auto",
                             target_lang: str = "tr") -> bytes:
        """
        Dönüştür ve çevir

        Args:
            pdf_bytes: PDF bayt verisi
            source_lang: Kaynak dil
            target_lang: Hedef dil

        Returns:
            bytes: Çevrilmiş DOCX
        """
        # Önce Word'e çevir
        docx_bytes = self.convert_with_pdf2docx(pdf_bytes)

        # Sonra çeviri yap
        return self._translate_docx(docx_bytes, source_lang, target_lang)

    def _translate_docx(self, docx_bytes: bytes, source_lang: str,
                       target_lang: str) -> bytes:
        """DOCX dosyasını çevir"""
        # DOCX'i yükle
        doc = Document(io.BytesIO(docx_bytes))

        # Her paragrafı çevir
        for para in doc.paragraphs:
            if para.text.strip():
                result = self.translator.translate(
                    para.text,
                    target_lang=target_lang,
                    source_lang=source_lang
                )

                if result.success:
                    # Metni güncelle
                    para.text = result.text

        # Tabloları çevir
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        result = self.translator.translate(
                            cell.text,
                            target_lang=target_lang,
                            source_lang=source_lang
                        )

                        if result.success:
                            cell.text = result.text

        # Sonucu bayt olarak döndür
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()


def convert_pdf_to_word(pdf_bytes: bytes, translate: bool = False,
                       source_lang: str = "auto", target_lang: str = "tr") -> bytes:
    """
    PDF'ten Word'e dönüşüm (kolay fonksiyon)

    Args:
        pdf_bytes: PDF bayt verisi
        translate: Çeviri yap
        source_lang: Kaynak dil
        target_lang: Hedef dil

    Returns:
        bytes: DOCX bayt verisi
    """
    converter = AdvancedPDFToWordConverter()

    if translate:
        return converter.convert_and_translate(pdf_bytes, source_lang, target_lang)
    else:
        return converter.convert_with_pdf2docx(pdf_bytes)
